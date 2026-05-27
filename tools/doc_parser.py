"""
doc_parser.py
─────────────
Parses uploaded files (.xlsx, .csv, .docx) into a clean plain-text string
that can be sent directly to the LLM for test-case generation.

Supported formats
-----------------
  .xlsx  → pandas reads every sheet; each cell value is tab-separated.
  .csv   → pandas reads with automatic delimiter detection.
  .docx  → python-docx extracts paragraph text + table cell text.
"""

from __future__ import annotations

import io
from typing import BinaryIO


# ─────────────────────────────────────────────────────────────
# Public API
# ─────────────────────────────────────────────────────────────

def parse_uploaded_file(uploaded_file) -> str:
    """
    Accept a Streamlit UploadedFile object (or any file-like with .name and .read())
    and return a single UTF-8 string with all extracted text.

    Raises ValueError for unsupported file types or parse failures.
    """
    name: str = uploaded_file.name.lower()
    raw_bytes: bytes = uploaded_file.read()

    try:
        if name.endswith(".xlsx") or name.endswith(".xls"):
            return _parse_excel(raw_bytes)
        elif name.endswith(".csv"):
            return _parse_csv(raw_bytes)
        elif name.endswith(".docx"):
            return _parse_docx(raw_bytes)
        else:
            raise ValueError(
                f"Unsupported file type: '{uploaded_file.name}'. "
                "Please upload a .xlsx, .xls, .csv, or .docx file."
            )
    except ValueError:
        raise  # Re-raise ValueError (unsupported type / explicit error)
    except ImportError as ie:
        raise ValueError(
            f"Thiếu thư viện cần thiết để đọc file này: {ie}. "
            "Hãy chạy: pip install openpyxl python-docx"
        ) from ie
    except Exception as exc:
        raise ValueError(
            f"Không thể đọc file '{uploaded_file.name}': {type(exc).__name__}: {exc}"
        ) from exc


# ─────────────────────────────────────────────────────────────
# Internal parsers
# ─────────────────────────────────────────────────────────────

def _parse_excel(raw_bytes: bytes) -> str:
    """Read all sheets from an Excel file and convert to readable text."""
    import pandas as pd

    # Explicitly require openpyxl engine so a missing/broken install surfaces
    # a clear ImportError instead of a cryptic crash.
    try:
        import openpyxl  # noqa: F401  -- ensure it's importable
    except ImportError:
        raise ImportError(
            "openpyxl is required to read .xlsx/.xls files. "
            "Run: pip install openpyxl"
        )

    buf = io.BytesIO(raw_bytes)
    xl = pd.ExcelFile(buf, engine="openpyxl")
    sections: list[str] = []

    for sheet_name in xl.sheet_names:
        df = xl.parse(sheet_name, header=None, dtype=str)
        df = df.fillna("")

        # Drop fully empty rows & columns to reduce noise
        df = df.loc[~(df == "").all(axis=1)]
        df = df.loc[:, ~(df == "").all(axis=0)]

        if df.empty:
            continue

        rows_text = "\n".join(
            "\t".join(str(cell) for cell in row)
            for row in df.itertuples(index=False, name=None)
        )
        sections.append(f"=== Sheet: {sheet_name} ===\n{rows_text}")

    if not sections:
        return "(File trống hoặc không có dữ liệu)"

    return "\n\n".join(sections)


def _parse_csv(raw_bytes: bytes) -> str:
    """Read a CSV file, auto-detect delimiter, and convert to readable text."""
    import pandas as pd

    # Try UTF-8 first, fall back to latin-1 for Vietnamese files exported from Windows
    for encoding in ("utf-8", "utf-8-sig", "latin-1"):
        try:
            buf = io.StringIO(raw_bytes.decode(encoding))
            df = pd.read_csv(buf, dtype=str, sep=None, engine="python")
            df = df.fillna("")
            break
        except (UnicodeDecodeError, pd.errors.ParserError):
            continue
    else:
        return "(Không thể đọc file CSV — thử chuyển sang UTF-8)"

    if df.empty:
        return "(File CSV trống)"

    header = "\t".join(str(c) for c in df.columns)
    rows = "\n".join(
        "\t".join(str(cell) for cell in row)
        for row in df.itertuples(index=False, name=None)
    )
    return f"{header}\n{rows}"


def _parse_docx(raw_bytes: bytes) -> str:
    """Extract all paragraphs and table cells from a .docx file."""
    try:
        from docx import Document  # python-docx
    except ImportError:
        return (
            "⚠️ Thư viện 'python-docx' chưa được cài đặt.\n"
            "Chạy: pip install python-docx"
        )

    buf = io.BytesIO(raw_bytes)
    doc = Document(buf)
    lines: list[str] = []

    # Paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            lines.append(text)

    # Tables
    for table_idx, table in enumerate(doc.tables, 1):
        lines.append(f"\n--- Bảng {table_idx} ---")
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            lines.append("\t".join(cells))

    if not lines:
        return "(File DOCX trống hoặc không có nội dung văn bản)"

    return "\n".join(lines)
