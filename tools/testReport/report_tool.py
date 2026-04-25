import os
import datetime
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pydantic import BaseModel, Field
from browser_use import Controller


# 1. Định nghĩa cấu trúc dữ liệu rõ ràng để Agent không bị nhầm lẫn
from pydantic import BaseModel, Field
from browser_use import Controller
from typing import List


# 1. Định nghĩa Sub-model để schema rõ ràng hơn, tránh lỗi OpenAI
class UIItem(BaseModel):
    button: str = Field(description="Tên của nút bấm hoặc link")
    status: str = Field(description="Trạng thái (hoạt động/không hoạt động)")


class SaveReportParams(BaseModel):
    title: str = Field(
        description="Tiêu đề báo cáo, bắt buộc phải là URL của trang web được test"
    )
    ui_items: List[UIItem] = Field(
        description="Danh sách kết quả các nút bấm đã test",
    )
    security_summary: str = Field(
        description="Kết quả kiểm tra bảo mật (web tĩnh/động và các lỗi nếu có)",
    )
    filename: str = Field(
        description="Tên file xuất ra, bắt buộc có đuôi .docx (vd: bao_cao_trang_chu.docx)"
    )


def register_report_tool(controller: Controller):
    @controller.registry.action(
        "Tạo file báo cáo tổng hợp chi tiết dạng Word (.docx) chứa cả dữ liệu giao diện và bảo mật",
        param_model=SaveReportParams,
    )
    async def save_report_action(params: SaveReportParams) -> str:
        try:
            # Tự động tạo thư mục reports nếu chưa có
            os.makedirs("reports", exist_ok=True)
            filepath = os.path.join("reports", params.filename)

            doc = Document()

            # --- HEADER ---
            heading = doc.add_heading(f"BÁO CÁO KIỂM THỬ: {params.title}", level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            heading.runs[0].font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

            timestamp = datetime.datetime.now().strftime("%d/%m/%Y %H:%M:%S")
            ts_para = doc.add_paragraph(f"📅 Thời gian thực hiện: {timestamp}")
            ts_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            ts_para.runs[0].font.size = Pt(10)
            ts_para.runs[0].italic = True

            doc.add_paragraph()  # Dòng trống

            # --- SECTION 1: GIAO DIỆN (UI) ---
            doc.add_heading("1. Đánh giá Giao diện (UI Buttons)", level=2)

            if params.ui_items:
                # Tạo bảng Word đẹp mắt cho dữ liệu UI thay vì in dòng text
                table = doc.add_table(rows=1, cols=2)
                table.style = "Light Shading Accent 1"  # Style có sẵn của Word

                # Header bảng
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = "Tên phần tử (Button/Link)"
                hdr_cells[1].text = "Trạng thái"

                # Điền dữ liệu
                for item in params.ui_items:
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(item.get("button", "N/A"))
                    row_cells[1].text = str(item.get("status", "N/A"))
            else:
                doc.add_paragraph(
                    "Không ghi nhận được phần tử tương tác nào trên trang này."
                )

            doc.add_paragraph()

            # --- SECTION 2: BẢO MẬT (SECURITY) ---
            doc.add_heading("2. Đánh giá Bảo mật (Security)", level=2)
            doc.add_paragraph(params.security_summary)

            # --- LƯU FILE ---
            doc.save(filepath)
            return f"✅ Report: Đã xuất báo cáo thành công tại '{filepath}'"

        except Exception as e:
            return f"❌ Lỗi nghiêm trọng khi tạo báo cáo Word: {str(e)}"
